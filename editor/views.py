from django.shortcuts import render, get_object_or_404, redirect
from django.http import HttpResponseForbidden, JsonResponse, HttpResponse
from django.contrib.auth.decorators import login_required
from django.views.decorators.csrf import csrf_exempt
from django.contrib.auth.models import User
from django.utils.timezone import now
from .models import Document, DocumentPermission, Comment, DocumentVersion
from .forms import DocumentForm
from docx import Document as DocxDocument
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import json
import logging
from django.utils.safestring import mark_safe

logger = logging.getLogger(__name__)

@login_required
def document_view(request, doc_id):
    # Get the document
    document = get_object_or_404(Document, id=doc_id)
    
    # Get all permissions for the document
    user_permissions = DocumentPermission.objects.filter(document=document)
    
    # If the user is the owner, they can view the permissions
    if request.user == document.owner:
        permissions_list = user_permissions
        return render(request, "editor/document.html", {
            "document": document,
            "can_view_permissions": True,
            "permissions_list": permissions_list,
        })
    
    # If the user is not the owner, check if they have access
    try:
        permission = DocumentPermission.objects.get(user=request.user, document=document)
        
        # If the user has "edit" access, they can view the permissions
        if permission.access_level == DocumentPermission.EDIT:
            permissions_list = user_permissions
            return render(request, "editor/document.html", {
                "document": document,
                "can_view_permissions": True,
                "permissions_list": permissions_list,
            })
        
        # If the user has "read" access, they cannot view permissions
        elif permission.access_level == DocumentPermission.READ:
            return render(request, "editor/document.html", {
                "document": document,
                "can_view_permissions": False,
            })
    
    except DocumentPermission.DoesNotExist:
        # If the user has no access, return a forbidden response
        return HttpResponseForbidden("You don't have permission to access this document.")
    
    return render(request, "editor/document.html", {
        "document": document,
        "permissions_list": permissions_list,
        "can_view_permissions": can_view_permissions,
    })


def home(request):
    return render(request, 'home.html')

@login_required
def dashboard(request):
    # Fetch documents owned by the user
    owned_documents = Document.objects.filter(owner=request.user)
    
    # Fetch documents shared with the user and annotate the access level
    shared_documents_with_permissions = DocumentPermission.objects.filter(user=request.user)
    
    shared_documents = []
    for permission in shared_documents_with_permissions:
        shared_documents.append({
            "document": permission.document,
            "access_level": permission.access_level
        })
    
    context = {
        'owned_documents': owned_documents,
        'shared_documents': shared_documents,
    }
    
    return render(request, 'editor/dashboard.html', context)

@login_required
def profile_view(request):
    return render(request, "editor/profile.html")

@login_required
def create_document(request):
    if request.method == "POST":
        form = DocumentForm(request.POST)
        if form.is_valid():
            document = form.save(commit=False)
            document.owner = request.user
            document.save()
            return redirect("edit_document", doc_id=document.id)
    else:
        form = DocumentForm()
    return render(request, "editor/create_document.html", {"form": form})

@login_required
def edit_document(request, doc_id):
    document = get_object_or_404(Document, id=doc_id)
    VERSION_THRESHOLD = 50  # Create a new version after 50 character changes

    # Check permissions
    if request.user != document.owner:
        try:
            permission = DocumentPermission.objects.get(user=request.user, document=document)
            if permission.access_level != DocumentPermission.EDIT:
                return HttpResponseForbidden("You don't have permission to edit this document.")
        except DocumentPermission.DoesNotExist:
            return HttpResponseForbidden("You don't have permission to edit this document.")

    # Handling POST request to save edits
    if request.method == "POST":
        new_content = request.POST.get('content', '')

        # Calculate the number of characters changed
        current_content = document.content or ""
        changed_characters = abs(len(new_content) - len(current_content))

        if request.method == "POST":
            new_content = request.POST.get('content', '')
            current_content = document.content or ""
        
            try:
                # Parse both contents as JSON to compare actual content
                new_content_json = json.loads(new_content)
                current_content_json = json.loads(current_content) if current_content else {"ops": []}
                
                # Compare the text content length
                new_text = ''.join(op.get('insert', '') for op in new_content_json.get('ops', []))
                current_text = ''.join(op.get('insert', '') for op in current_content_json.get('ops', []))
                
                changed_characters = abs(len(new_text) - len(current_text))
                
                if changed_characters >= VERSION_THRESHOLD:
                    # Create new version
                    DocumentVersion.objects.create(
                        document=document,
                        version_number=document.versions.count() + 1,
                        content=current_content,
                        modified_at=now(),
                        modified_by=request.user
                    )
                    print(f"New version created for document: {document.title}")
                
                # Update document content
                document.content = new_content
                document.save()
                
            except json.JSONDecodeError:
                # Handle invalid JSON gracefully
                print("Error parsing document content JSON")
                document.content = new_content
                document.save()

        # Update the document with new content
        document.content = new_content
        document.save()
        print(f"Document '{document.title}' updated with new content.")

    # Fetch permissions list for display
    permissions_list = DocumentPermission.objects.filter(document=document)
    comments = document.comments.all()  # Fetch comments for the document

    return render(request, "editor/edit_document.html", {
        "document": document,
        "permissions_list": permissions_list,
        "comments": comments,
    })


@login_required
def share_document(request, doc_id):
    document = get_object_or_404(Document, id=doc_id, owner=request.user)

    if request.method == "POST":
        user_identifier = request.POST.get('user_identifier')
        access_level = request.POST.get('access_level')  # Get access level from form

        # Check if the user exists by username or email
        try:
            user = User.objects.get(username=user_identifier) if '@' not in user_identifier else User.objects.get(email=user_identifier)
            # Update or create the permission
            permission, created = DocumentPermission.objects.update_or_create(
                user=user,
                document=document,
                defaults={'access_level': access_level},
            )
            return redirect('edit_document', doc_id=doc_id)
        except User.DoesNotExist:
            return render(request, 'editor/share_document.html', {
                'error': "User not found",
                'document': document
            })

    return render(request, 'editor/share_document.html', {'document': document})

@login_required
def manage_access(request, doc_id):
    document = get_object_or_404(Document, id=doc_id, owner=request.user)
    permissions = DocumentPermission.objects.filter(document=document)

    if request.method == "POST":
        user_id = request.POST.get("user_id")
        new_access_level = request.POST.get("access_level")
        permission = get_object_or_404(DocumentPermission, id=user_id, document=document)
        permission.access_level = new_access_level
        permission.save()

    return render(request, "editor/manage_access.html", {"document": document, "permissions": permissions})


@login_required
def read_document(request, doc_id):
    document = get_object_or_404(Document, id=doc_id)
    
    # Allow access if the user is the owner
    if request.user == document.owner:
        return render(request, "editor/read_document.html", {"document": document})
    
    # Otherwise, check the permission level
    try:
        permission = DocumentPermission.objects.get(user=request.user, document=document)
        if permission.access_level == DocumentPermission.READ:
            return render(request, "editor/read_document.html", {"document": document})
        else:
            return HttpResponseForbidden("You don't have permission to view this document.")
    except DocumentPermission.DoesNotExist:
        return HttpResponseForbidden("You don't have permission to view this document.")

@csrf_exempt
@login_required
def add_comment(request, doc_id):
    if request.method == "POST":
        try:
            document = Document.objects.get(pk=doc_id)
            data = json.loads(request.body)

            content = data.get("content")
            range_start = data.get("range_start")
            range_end = data.get("range_end")

            print("Received comment data:", content, range_start, range_end)  # Debug log

            comment = Comment.objects.create(
                document=document,
                user=request.user,
                content=content,
                range_start=range_start,
                range_end=range_end,
            )

            print("Comment saved:", comment)  # Debug log

            return JsonResponse({
                "success": True,
                "comment": {
                    "user": comment.user.username,
                    "content": comment.content,
                    "created_at": comment.created_at.strftime("%Y-%m-%d %H:%M:%S"),
                }
            })
        except Document.DoesNotExist:
            print("Document not found")  # Debug log
            return JsonResponse({"success": False, "error": "Document not found"}, status=404)
    return JsonResponse({"success": False, "error": "Invalid request method"}, status=400)

@csrf_exempt
def save_document(request, doc_id):
    if request.method == "POST":
        try:
            document = Document.objects.get(id=doc_id)
            data = json.loads(request.body)
            document.content = data.get("content", "")
            document.save()
            return JsonResponse({"success": True})
        except Document.DoesNotExist:
            return JsonResponse({"error": "Document not found"}, status=404)
    return JsonResponse({"error": "Invalid request method"}, status=405)

@login_required
def download_document(request, doc_id):
    document = get_object_or_404(Document, id=doc_id)
    
    # Permission check
    if request.user != document.owner:
        try:
            permission = DocumentPermission.objects.get(user=request.user, document=document)
            if permission.access_level not in [DocumentPermission.READ, DocumentPermission.EDIT]:
                return JsonResponse({'error': 'Permission denied'}, status=403)
        except DocumentPermission.DoesNotExist:
            return JsonResponse({'error': 'Permission denied'}, status=403)
    
    # Create a new Word document
    docx = DocxDocument()
    
    # Add title
    title = document.title
    docx.add_heading(title, level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Process content stored in the database
    try:
        content = json.loads(document.content)
        if content and 'ops' in content:
            for op in content['ops']:
                insert = op.get('insert', '')
                attributes = op.get('attributes', {})
                
                # Add a paragraph
                paragraph = docx.add_paragraph()
                run = paragraph.add_run(insert)
                
                # Apply formatting
                if attributes:
                    if attributes.get('bold'):
                        run.bold = True
                    if attributes.get('italic'):
                        run.italic = True
                    if attributes.get('underline'):
                        run.underline = True
                    if attributes.get('size'):
                        size_map = {
                            'small': Pt(8),
                            'large': Pt(14),
                            'huge': Pt(18)
                        }
                        run.font.size = size_map.get(attributes['size'], Pt(11))
    except json.JSONDecodeError:
        return JsonResponse({'error': 'Invalid content format'}, status=400)
    
    # Save the document to an in-memory file
    doc_stream = io.BytesIO()
    docx.save(doc_stream)
    doc_stream.seek(0)
    
    # Create response for downloading the document
    response = HttpResponse(
        doc_stream.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename="{document.title}.docx"'
    
    return response

@login_required
def view_versions(request, doc_id):
    document = get_object_or_404(Document, id=doc_id)
    
    # Permission check
    if request.user != document.owner:
        try:
            permission = DocumentPermission.objects.get(user=request.user, document=document)
            if permission.access_level not in [DocumentPermission.READ, DocumentPermission.EDIT]:
                return HttpResponseForbidden("You don't have permission to view versions of this document.")
        except DocumentPermission.DoesNotExist:
            return HttpResponseForbidden("You don't have permission to view versions of this document.")

    # Get versions with proper ordering and annotation
    versions = (DocumentVersion.objects
               .filter(document=document)
               .select_related('modified_by')  # Optimize query
               .order_by('-version_number'))
    
    return render(request, "editor/view_versions.html", {
        "document": document,
        "versions": versions,
    })

@login_required
def get_document_versions(request, doc_id):
    document = get_object_or_404(Document, id=doc_id)
    
    # Permission check
    if request.user != document.owner:
        try:
            permission = DocumentPermission.objects.get(user=request.user, document=document)
            if permission.access_level not in [DocumentPermission.READ, DocumentPermission.EDIT]:
                return JsonResponse({'error': 'Permission denied'}, status=403)
        except DocumentPermission.DoesNotExist:
            return JsonResponse({'error': 'Permission denied'}, status=403)
    
    try:
        versions = document.versions.all().order_by('-version_number')
        versions_data = []
        
        for version in versions:
            try:
                content = json.loads(version.content) if version.content else {"ops": []}
                versions_data.append({
                    'id': version.id,
                    'version_number': version.version_number,
                    'content': version.content,
                    'modified_at': version.modified_at.isoformat(),
                    'modified_by': version.modified_by.username,
                    'text_content': ''.join(op.get('insert', '') for op in content.get('ops', []))
                })
            except json.JSONDecodeError:
                logger.error(f"Failed to parse content for version {version.id}")
                continue
                
        return JsonResponse({'versions': versions_data})
        
    except Exception as e:
        logger.error(f"Error in get_document_versions: {str(e)}")
        return JsonResponse({'error': 'Failed to fetch versions'}, status=500)

@login_required
def restore_version(request, doc_id, version_id):
    if request.method != 'POST':
        return JsonResponse({'error': 'Method not allowed'}, status=405)
        
    document = get_object_or_404(Document, id=doc_id)
    version = get_object_or_404(DocumentVersion, id=version_id, document=document)
    
    # Check permissions
    if request.user != document.owner:
        try:
            permission = DocumentPermission.objects.get(user=request.user, document=document)
            if permission.access_level != DocumentPermission.EDIT:
                return JsonResponse({'error': 'Permission denied'}, status=403)
        except DocumentPermission.DoesNotExist:
            return JsonResponse({'error': 'Permission denied'}, status=403)
    
    # Create a new version with the current content before restoring
    DocumentVersion.objects.create(
        document=document,
        version_number=document.versions.count() + 1,
        content=document.content,
        modified_by=request.user
    )
    
    # Restore the selected version
    document.content = version.content
    document.save()
    
    return JsonResponse({'success': True})

@login_required
def view_version_content(request, doc_id, version_id):
    document = get_object_or_404(Document, id=doc_id)
    version = get_object_or_404(DocumentVersion, id=version_id, document=document)

    # Check permissions
    if request.user != document.owner:
        try:
            permission = DocumentPermission.objects.get(user=request.user, document=document)
            if permission.access_level not in [DocumentPermission.READ, DocumentPermission.EDIT]:
                return HttpResponseForbidden("You don't have permission to view this version.")
        except DocumentPermission.DoesNotExist:
            return HttpResponseForbidden("You don't have permission to view this version.")

    # Parse the version content (which is in Quill Delta format)
    try:
        delta = json.loads(version.content)
        parsed_content = parse_quill_delta(delta)
    except json.JSONDecodeError:
        parsed_content = "Error parsing content"

    return render(request, "editor/view_version_content.html", {
        "document": document,
        "version": version,
        "parsed_content": mark_safe(parsed_content),  # Use mark_safe to render HTML safely
    })

def parse_quill_delta(delta):
    html_content = ""
    for op in delta.get('ops', []):
        text = op.get('insert', '')
        attributes = op.get('attributes', {})

        if 'bold' in attributes:
            text = f"<strong>{text}</strong>"
        if 'italic' in attributes:
            text = f"<em>{text}</em>"
        if 'underline' in attributes:
            text = f"<u>{text}</u>"
        if 'size' in attributes:
            size_map = {
                'small': '0.8em',
                'large': '1.5em',
                'huge': '2em'
            }
            size = size_map.get(attributes['size'], '1em')
            text = f"<span style='font-size: {size};'>{text}</span>"

        html_content += text.replace('\n', '<br>')

    return html_content

@login_required
def view_versions(request, doc_id):
    document = get_object_or_404(Document, id=doc_id)

    # Check if the user has permission to view the document
    if request.user != document.owner:
        try:
            permission = DocumentPermission.objects.get(user=request.user, document=document)
            if permission.access_level not in [DocumentPermission.READ, DocumentPermission.EDIT]:
                return HttpResponseForbidden("You don't have permission to view versions of this document.")
        except DocumentPermission.DoesNotExist:
            return HttpResponseForbidden("You don't have permission to view versions of this document.")

    # Get all versions of the document
    versions = document.versions.all().order_by('-version_number')  # Order by descending version number
    print("Versions retrieved:", versions)  # Debugging line to see if versions are being retrieved

    return render(request, "editor/view_versions.html", {
        "document": document,
        "versions": versions,
    })


